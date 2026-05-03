"""Генерує базові DOCX-шаблони для сценарію «невиплата зарплати».

Це ЗАГЛУШКИ — мінімально юридично коректні шаблони, які будуть замінені
на робочі шаблони від Дмитра Глушка перед публічним релізом.

Запуск:
    python scripts/build_templates.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
TEMPLATES_DIR = ROOT / "templates"

DISCLAIMER = (
    "_____________________________________________________________________\n"
    "Цей документ згенеровано автоматично сервісом PravoHelpAI станом на "
    "{{ today }}. Перед поданням рекомендуємо перевірку фахового юриста. "
    "Шаблон є типовим і потребує адаптації під обставини конкретної справи."
)


def _setup(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)


def _header(doc: Document, lines: list[str], align: int = WD_ALIGN_PARAGRAPH.RIGHT) -> None:
    for line in lines:
        p = doc.add_paragraph(line)
        p.alignment = align


def _title(doc: Document, text: str) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()


def _para(doc: Document, text: str, *, indent: bool = True, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.bold = bold


def _signature(doc: Document) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph("«{{ today }}»")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2 = doc.add_paragraph("________________ / {{ user_name }} /")
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _disclaimer(doc: Document) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(DISCLAIMER)
    run.italic = True
    run.font.size = Pt(9)


# ============================================================
# 1. Претензія роботодавцю
# ============================================================
def build_employer_claim() -> None:
    doc = Document()
    _setup(doc)

    _header(doc, [
        "Керівнику {{ employer_name }}",
        "{{ employer_address }}",
        "ЄДРПОУ: {{ employer_edrpou_or_dash }}",
        "",
        "Від {{ user_name }}",
        "ІПН: {{ user_tax_id }}",
        "Адреса: {{ user_address }}",
        "Тел.: {{ user_phone }}",
    ])

    _title(doc, "ДОСУДОВА ПРЕТЕНЗІЯ")
    _para(doc, "про виплату заборгованості із заробітної плати")

    _para(
        doc,
        "Я, {{ user_name }}, перебуваю у трудових відносинах із {{ employer_name }}. "
        "Станом на {{ today }} мені не виплачено заробітну плату за період "
        "з {{ period_from_text }} по {{ period_to_text }} у розмірі {{ amount_text }}.",
    )
    _para(
        doc,
        "Останню виплату заробітної плати було здійснено {{ last_payment_date }}.",
    )
    _para(
        doc,
        "Відповідно до статті 115 Кодексу законів про працю України заробітна плата "
        "виплачується працівникам регулярно в робочі дні у строки, встановлені колективним "
        "договором або нормативним актом роботодавця, але не рідше двох разів на місяць "
        "через проміжок часу, що не перевищує шістнадцяти календарних днів.",
    )
    _para(
        doc,
        "Згідно зі статтею 116 КЗпП України, при припиненні трудового договору виплата всіх "
        "сум, що належать працівникові, провадиться в день звільнення. Стаття 117 КЗпП України "
        "передбачає відповідальність роботодавця за затримку розрахунку у формі середнього "
        "заробітку працівника за весь час затримки по день фактичного розрахунку.",
    )

    _para(doc, "На підставі викладеного, керуючись ст.ст. 115, 116, 117 КЗпП України,",
          indent=False, bold=True)

    _para(doc, "ВИМАГАЮ:", indent=False, bold=True)
    _para(
        doc,
        "1. У строк до 7 (семи) календарних днів з дня отримання цієї претензії виплатити "
        "заборгованість із заробітної плати у розмірі {{ amount_text }}.",
    )
    _para(
        doc,
        "2. У разі затримки виплати — виплатити середній заробіток за весь час затримки "
        "відповідно до ст. 117 КЗпП України.",
    )
    _para(
        doc,
        "3. Письмово повідомити мене про результати розгляду цієї претензії за вказаною вище адресою.",
    )

    _para(
        doc,
        "У разі залишення цієї претензії без задоволення, я буду змушений(а) звернутися до "
        "Державної служби України з питань праці та/або до суду з відповідним позовом, з покладенням "
        "на роботодавця всіх судових витрат.",
    )

    _signature(doc)
    _disclaimer(doc)

    out = TEMPLATES_DIR / "salary_claim_employer.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")


# ============================================================
# 2. Скарга до Держпраці
# ============================================================
def build_labor_office_claim() -> None:
    doc = Document()
    _setup(doc)

    _header(doc, [
        "До Державної служби України з питань праці",
        "(територіальне управління Держпраці",
        "за місцем реєстрації роботодавця)",
        "",
        "Від {{ user_name }}",
        "ІПН: {{ user_tax_id }}",
        "Адреса: {{ user_address }}",
        "Тел.: {{ user_phone }}",
    ])

    _title(doc, "ЗАЯВА")
    _para(doc, "про порушення законодавства про працю — невиплату заробітної плати", indent=True)

    _para(
        doc,
        "Я, {{ user_name }}, перебуваю у трудових відносинах із {{ employer_name }} "
        "(ЄДРПОУ {{ employer_edrpou_or_dash }}, адреса: {{ employer_address }}).",
    )
    _para(
        doc,
        "Станом на {{ today }} мені не виплачено заробітну плату за період "
        "з {{ period_from_text }} по {{ period_to_text }} у розмірі {{ amount_text }}. "
        "Остання виплата зарплати була проведена {{ last_payment_date }}.",
    )
    _para(
        doc,
        "Зазначені дії роботодавця порушують ст.ст. 115, 116 КЗпП України, що визначають "
        "обовʼязок роботодавця регулярно і своєчасно виплачувати заробітну плату.",
    )

    _para(doc, "Керуючись ст. 259 КЗпП України та Положенням про Державну службу України з питань праці,",
          indent=False, bold=True)
    _para(doc, "ПРОШУ:", indent=False, bold=True)
    _para(
        doc,
        "1. Провести перевірку додержання роботодавцем — {{ employer_name }} — "
        "законодавства про оплату праці.",
    )
    _para(
        doc,
        "2. Вжити заходів реагування у разі виявлення порушень, у тому числі застосувати "
        "до роботодавця штрафні санкції згідно зі ст. 265 КЗпП України.",
    )
    _para(
        doc,
        "3. Письмово повідомити мене про результати розгляду цієї заяви за вказаною вище адресою.",
    )

    _signature(doc)
    _disclaimer(doc)

    out = TEMPLATES_DIR / "salary_claim_labor_office.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")


# ============================================================
# 3. Позовна заява до суду
# ============================================================
def build_court_claim() -> None:
    doc = Document()
    _setup(doc)

    _header(doc, [
        "До {{ court_name_or_placeholder }}",
        "(районний суд за місцем реєстрації",
        "роботодавця або позивача)",
        "",
        "Позивач: {{ user_name }}",
        "ІПН: {{ user_tax_id }}",
        "Адреса: {{ user_address }}",
        "Тел.: {{ user_phone }}",
        "",
        "Відповідач: {{ employer_name }}",
        "ЄДРПОУ: {{ employer_edrpou_or_dash }}",
        "Адреса: {{ employer_address }}",
        "",
        "Ціна позову: {{ amount_text }}",
        "Судовий збір: позивач звільнений від сплати",
        "відповідно до п. 1 ч. 1 ст. 5 ЗУ «Про судовий збір»",
    ])

    _title(doc, "ПОЗОВНА ЗАЯВА")
    _para(doc, "про стягнення заборгованості із заробітної плати", indent=True)

    _para(
        doc,
        "Я, {{ user_name }}, перебуваю у трудових відносинах із Відповідачем — "
        "{{ employer_name }}. Між Позивачем та Відповідачем укладено трудовий договір, "
        "згідно з яким Позивач виконує трудові обовʼязки.",
    )
    _para(
        doc,
        "Станом на {{ today }} Відповідач не виплатив Позивачу заробітну плату за період "
        "з {{ period_from_text }} по {{ period_to_text }} у розмірі {{ amount_text }}. "
        "Остання виплата заробітної плати — {{ last_payment_date }}.",
    )
    _para(
        doc,
        "Відповідач письмових пояснень щодо причин невиплати не надав, заборгованість не погасив.",
    )

    _para(
        doc,
        "Згідно зі ст. 115 КЗпП України, заробітна плата виплачується працівникам регулярно в "
        "робочі дні. Стаття 116 КЗпП передбачає виплату всіх належних сум при припиненні "
        "трудового договору в день звільнення. Згідно зі ст. 117 КЗпП України, при затримці "
        "розрахунку Відповідач зобовʼязаний виплатити середній заробіток за весь час затримки.",
    )
    _para(
        doc,
        "Відповідно до п. 1 ч. 1 ст. 5 ЗУ «Про судовий збір», позивач у справах про стягнення "
        "заробітної плати звільнений від сплати судового збору.",
    )

    _para(doc, "На підставі викладеного, керуючись ст.ст. 115, 116, 117 КЗпП України, ст.ст. 4, 175 ЦПК України,",
          indent=False, bold=True)
    _para(doc, "ПРОШУ:", indent=False, bold=True)
    _para(
        doc,
        "1. Стягнути з {{ employer_name }} (ЄДРПОУ {{ employer_edrpou_or_dash }}) на користь "
        "{{ user_name }} заборгованість із заробітної плати у розмірі {{ amount_text }}.",
    )
    _para(
        doc,
        "2. Стягнути з Відповідача середній заробіток за весь час затримки розрахунку "
        "відповідно до ст. 117 КЗпП України.",
    )
    _para(
        doc,
        "3. Стягнути з Відповідача судові витрати, понесені Позивачем у справі.",
    )

    _para(doc, "Додатки:", indent=False, bold=True)
    _para(doc, "1. Копія паспорта Позивача — на 1 арк.", indent=False)
    _para(doc, "2. Копія РНОКПП Позивача — на 1 арк.", indent=False)
    _para(doc, "3. Копія трудового договору / витяг з трудової книжки — на ___ арк.", indent=False)
    _para(doc, "4. Розрахунок заборгованості — на 1 арк.", indent=False)
    _para(doc, "5. Копія цієї позовної заяви для Відповідача — на ___ арк.", indent=False)

    _signature(doc)
    _disclaimer(doc)

    out = TEMPLATES_DIR / "salary_court_claim.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")


def main() -> None:
    TEMPLATES_DIR.mkdir(exist_ok=True)
    print(f"Генерую шаблони у {TEMPLATES_DIR}/")
    build_employer_claim()
    build_labor_office_claim()
    build_court_claim()
    print("Готово.")


if __name__ == "__main__":
    main()
